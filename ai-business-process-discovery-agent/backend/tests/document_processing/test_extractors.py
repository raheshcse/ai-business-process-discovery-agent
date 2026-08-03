from pathlib import Path

import pandas as pd
import pymupdf
from docx import Document as WordDocument

from app.document_processing.extractors import (
    CsvExtractor,
    DocxExtractor,
    PdfExtractor,
    TxtExtractor,
    XlsxExtractor,
)


def test_pdf_extractor_returns_text_and_page_count(tmp_path: Path) -> None:
    path = tmp_path / "sample.pdf"
    document = pymupdf.open()
    first_page = document.new_page()
    first_page.insert_text((72, 72), "First page")
    second_page = document.new_page()
    second_page.insert_text((72, 72), "Second page")
    document.save(path)
    document.close()

    result = PdfExtractor().extract(path)

    assert result.page_count == 2
    assert "First page" in result.text
    assert "Second page" in result.text


def test_docx_extractor_returns_paragraphs_and_table_cells(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = WordDocument()
    document.add_paragraph("Process overview")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Owner"
    table.cell(0, 1).text = "Operations"
    document.save(path)

    result = DocxExtractor().extract(path)

    assert "Process overview" in result.text
    assert "Owner | Operations" in result.text
    assert result.page_count is None


def test_txt_extractor_reads_utf8_and_removes_bom(tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"
    path.write_text("\ufeffDiscovery notes", encoding="utf-8")

    result = TxtExtractor().extract(path)

    assert result.text == "Discovery notes"


def test_csv_extractor_preserves_headers_and_values(tmp_path: Path) -> None:
    path = tmp_path / "sample.csv"
    path.write_text("step,owner\nReview,Finance\n", encoding="utf-8")

    result = CsvExtractor().extract(path)

    assert result.text == "step,owner\nReview,Finance\n"


def test_xlsx_extractor_includes_each_sheet(tmp_path: Path) -> None:
    path = tmp_path / "sample.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame({"step": ["Review"], "owner": ["Finance"]}).to_excel(
            writer, sheet_name="Current State", index=False
        )
        pd.DataFrame({"risk": ["Delay"]}).to_excel(
            writer, sheet_name="Risks", index=False
        )

    result = XlsxExtractor().extract(path)

    assert "Sheet: Current State" in result.text
    assert "Review,Finance" in result.text
    assert "Sheet: Risks" in result.text
    assert "Delay" in result.text
