from pathlib import Path

import pandas as pd
import pymupdf
from docx import Document as WordDocument

from app.document_processing.base import DocumentExtractor
from app.document_processing.models import DocumentType, RawExtraction


class PdfExtractor(DocumentExtractor):
    document_type = DocumentType.PDF
    supported_extensions = frozenset({".pdf"})

    def extract(self, file_path: Path) -> RawExtraction:
        self.validate_file(file_path)
        with pymupdf.open(file_path) as document:
            return RawExtraction(
                text="\n\n".join(page.get_text("text") for page in document),
                page_count=document.page_count,
            )


class DocxExtractor(DocumentExtractor):
    document_type = DocumentType.DOCX
    supported_extensions = frozenset({".docx"})

    def extract(self, file_path: Path) -> RawExtraction:
        self.validate_file(file_path)
        document = WordDocument(file_path)
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            parts.extend(
                " | ".join(cell.text for cell in row.cells) for row in table.rows
            )
        return RawExtraction(text="\n".join(parts))


class TxtExtractor(DocumentExtractor):
    document_type = DocumentType.TXT
    supported_extensions = frozenset({".txt"})

    def extract(self, file_path: Path) -> RawExtraction:
        self.validate_file(file_path)
        return RawExtraction(text=file_path.read_text(encoding="utf-8-sig"))


class CsvExtractor(DocumentExtractor):
    document_type = DocumentType.CSV
    supported_extensions = frozenset({".csv"})

    def extract(self, file_path: Path) -> RawExtraction:
        self.validate_file(file_path)
        frame = pd.read_csv(file_path, dtype=str, keep_default_na=False)
        return RawExtraction(text=frame.to_csv(index=False, lineterminator="\n"))


class XlsxExtractor(DocumentExtractor):
    document_type = DocumentType.XLSX
    supported_extensions = frozenset({".xlsx"})

    def extract(self, file_path: Path) -> RawExtraction:
        self.validate_file(file_path)
        sheets = pd.read_excel(
            file_path,
            sheet_name=None,
            dtype=str,
            keep_default_na=False,
            engine="openpyxl",
        )
        parts: list[str] = []
        for sheet_name, frame in sheets.items():
            parts.append(f"Sheet: {sheet_name}")
            parts.append(frame.to_csv(index=False, lineterminator="\n").rstrip())
        return RawExtraction(text="\n\n".join(parts))
